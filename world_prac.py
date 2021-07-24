import docx

doc = docx.Document('设备维修协议.docx')
# print(len(doc.paragraphs))
# for n in range(28):
#     print(doc.paragraphs[n].text)

print(doc.tables[0].cell(0,0).text)