import openpyxl
import docx

class Asset:
    """固定资产的类"""

    def __init__(self, asset_info, order_number):
        self.form_name = '医疗设备维修及配件采购申请新.xlsx'
        self.sheet_name = 'Sheet1'
        self.contract_name = '设备维修协议.docx'
        self.receiving_name = '设备验收报告.docx'
        self.asset_info = asset_info
        self.order_number = order_number
        self.maintain_order_1 = asset_info['设备维修'][order_number]

    def creat_form(self):
        """创建维修及配件申请单"""
        wb = openpyxl.load_workbook(self.form_name)
        sheet = wb[self.sheet_name]

        sheet['B2'] = self.asset_info['使用部门']
        sheet['F2'] = self.maintain_order_1['报修内容']['报修时间']
        sheet['B3'] = self.asset_info['固定资产名称']
        sheet['F3'] = self.asset_info['固定资产编号']
        sheet['B4'] = self.asset_info['原值']
        sheet['F4'] = self.asset_info['标准生产厂家']
        sheet['B5'] = self.asset_info['标准型号']
        sheet['F5'] = self.asset_info['标准经销商']
        sheet['B6'] = self.asset_info['序列号']
        sheet['B7'] = self.maintain_order_1['报修内容']['故障描述']

        parts_info = self.maintain_order_1['parts_info']
        row_n = 10
        for key in parts_info:
            sheet.cell(row=row_n, column=1).value = parts_info[key]['配件名称']
            sheet.cell(row=row_n, column=2).value = parts_info[key]['配件型号']
            sheet.cell(row=row_n, column=4).value = parts_info[key]['配件价格']
            sheet.cell(row=row_n, column=5).value = parts_info[key]['配件数量']
            sheet.cell(row=row_n, column=6).value = \
sheet.cell(row=row_n, column=4).value * sheet.cell(row=row_n, column=5).value
            row_n += 1

        sheet['F15'] = '=SUM(F10:F14)'

        sheet['B17'] = self.maintain_order_1['维修代理商']['维修代理商']
        sheet['F17'] = self.maintain_order_1['维修代理商']['联系方式']

        form_name_n = f"{self.asset_info['固定资产编号']}_{self.order_number}.xlsx"
        form_name_n = f'医疗设备维修及配件采购申请新_{form_name_n}'
        wb.save(form_name_n)
        print('维修申请单已经创建！')

    def creat_contract(self):
        """创建维修合同"""
        doc = docx.Document(self.contract_name)
        # 甲方设备信息
        table_1 = doc.tables[0]
        table_1.cell(0,0).text += self.asset_info['固定资产编号']
        table_1.cell(0,1).text += self.asset_info['标准生产厂家']
        table_1.cell(1,0).text += self.asset_info['固定资产名称']
        table_1.cell(1,1).text += self.asset_info['标准经销商']
        table_1.cell(2,0).text += self.asset_info['标准型号']
        table_1.cell(2,1).text += self.asset_info['序列号']
        table_1.cell(3,0).text += self.asset_info['开始使用日期']
        table_1.cell(4,0).text += self.asset_info['使用部门']
        table_1.cell(4,1).text += self.asset_info['存放地点']

        # 乙方配件信息
        # table_2 = doc.tables[1]

        # 故障描述
        para_obj = doc.paragraphs[7]
        para_obj.add_run(self.asset_info['设备维修']
            [self.order_number]['报修内容']['故障描述'])

        form_name_n = f"{self.asset_info['固定资产编号']}_{self.order_number}.docx"
        form_name_n = f'设备维修协议_{form_name_n}'
        doc.save(form_name_n)
        print('设备维修协议已经创建！')

    def creat_receiving_report(self):
        """创建验收报告"""
        doc = docx.Document(self.receiving_name)
        # 设备基本信息
        table = doc.tables[0]
        table.cell(1,0).text += self.asset_info['固定资产编号']
        table.cell(1,3).text += self.asset_info['标准生产厂家']
        table.cell(2,0).text += self.asset_info['固定资产名称']
        table.cell(2,3).text += self.asset_info['标准经销商']
        table.cell(3,0).text += self.asset_info['标准型号']
        table.cell(3,3).text += self.asset_info['序列号']
        table.cell(4,0).text += self.asset_info['开始使用日期']
        table.cell(5,0).text += self.asset_info['使用部门']
        table.cell(5,3).text += self.asset_info['存放地点']
        # 维修记录
        maintain_info = self.asset_info['设备维修'][self.order_number]
        table.cell(7,0).text += maintain_info['报修内容']['报修时间']
        table.cell(8,0).text += maintain_info['报修内容']['报修人']
        table.cell(8,3).text += maintain_info['报修内容']['报修人联系方式']
        table.cell(9,0).add_paragraph(maintain_info['报修内容']['故障描述'])

        form_name_n = f"{self.asset_info['固定资产编号']}_{self.order_number}.docx"
        form_name_n = f'设备验收报告_{form_name_n}'
        doc.save(form_name_n)
        print('设备验收报告已经创建！')


if __name__ == '__main__':
    """创建一个固定资产的类，并测试程序"""
    asset_info = {'使用年限': 60,
                    '使用部门': '小儿骨科病房',
                    '原值': 4450,
                    '合同号': 'LNZC20180501314',
                    '固定资产名称': '输液泵',
                    '固定资产编号': 'NS-20191009089',
                    '国别': '中国',
                    '存放地点': '南湖2号楼12层',
                    '实际地点编码': '010212',
                    '序列号': 'XA00180403A039',
                    '开始使用日期': '2019.10.10',
                    '数量': 1,
                    '标准型号': 'ZNB-XAII',
                    '标准存放地点': '南湖小儿骨科病房',
                    '标准生产厂家': '北京科力建元公司',
                    '标准经销商': '沈阳柳达公司',
                    '设备维修': {
                        '20210721001': {
                            '报修内容': {'故障描述': '测试', 
                                 '报修时间': '20210721',
                                 '报修人': 'repair_man',
                                 '报修人联系方式': 'repair_man_tel',},
                            '配件信息1':{'配件名称': 'part_name',
                                '配件型号': 'part_type',
                                '配件价格': 'part_price',
                                '配件数量': 'part_number',},
                            '维修代理商': {'维修代理商': 'agency_name',
                                '联系方式': 'agency_tel',},
                            '维修信息': {'维修内容': 'maintain_details',
                                '维修结束时间': '20210721',
                                '维修工程师': 'maintain_engineer',
                                '工程师联系方式': 'engineer_tel',},
                            },
                    },
                    '负责人': '胡晓云',
                    '负责工程师': '李凡'}
    asset_ex = Asset(asset_info, '20210721001')
    # asset_ex.creat_contract()
    asset_ex.creat_receiving_report()
